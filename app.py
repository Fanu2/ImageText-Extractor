import streamlit as st
from PIL import Image
import pytesseract
import pandas as pd
from io import BytesIO
from docx import Document

# 🎯 App title and description
st.set_page_config(page_title="Tesseract OCR App", layout="centered")
st.title("📝 Tesseract OCR Text Recognition")
st.markdown("Upload an image to extract text using Tesseract OCR.")

# 📤 Image upload
uploaded_file = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])

if uploaded_file:
    try:
        # 🖼️ Load and display image
        image = Image.open(uploaded_file).convert("RGB")
        st.image(image, caption="Uploaded Image", use_column_width=True)

        # 🔍 OCR
        with st.spinner("Running OCR..."):import streamlit as st
from PIL import Image
import pytesseract
import pandas as pd
from io import BytesIO
from docx import Document

# 🎯 App title and description
st.set_page_config(page_title="Tesseract OCR App", layout="centered")
st.title("📝 Tesseract OCR Text Recognition")
st.markdown("Upload an image to extract text using Tesseract OCR.")

# 📤 Image upload
uploaded_file = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])

if uploaded_file:
    try:
        # 🖼️ Load and display image
        image = Image.open(uploaded_file).convert("RGB")
        st.image(image, caption="Uploaded Image", use_column_width=True)

        # 🔍 OCR
        with st.spinner("Running OCR..."):
            extracted_text = pytesseract.image_to_string(image)

        # 📄 Display results
        if extracted_text.strip():
            st.subheader("📌 Extracted Text:")
            st.text_area("Text Output", extracted_text, height=300)

            # 🔄 Prepare exportable data
            extracted_lines = [{"Text": line.strip()} for line in extracted_text.splitlines() if line.strip()]
            df = pd.DataFrame(extracted_lines)

            # 📥 Export as CSV
            csv = df.to_csv(index=False).encode("utf-8")
            st.download_button("📥 Download CSV", csv, "ocr_output.csv", "text/csv")

            # 📥 Export as Excel
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                df.to_excel(writer, index=False, sheet_name="OCR Text")
            st.download_button("📥 Download Excel", excel_buffer.getvalue(), "ocr_output.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

            # 📥 Export as DOCX
            doc = Document()
            doc.add_heading("Extracted OCR Text", level=1)
            for row in extracted_lines:
                doc.add_paragraph(row["Text"])
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            st.download_button("📥 Download DOCX", doc_buffer.getvalue(), "ocr_output.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        else:
            st.warning("No text found in the image.")

    except Exception as e:
        st.error(f"⚠️ OCR failed: {e}")
else:
    st.info("Please upload an image to begin.")

            extracted_text = pytesseract.image_to_string(image)

        # 📄 Display results
        if extracted_text.strip():
            st.subheader("📌 Extracted Text:")
            st.text_area("Text Output", extracted_text, height=300)
        else:
            st.warning("No text found in the image.")

    except Exception as e:
        st.error(f"⚠️ OCR failed: {e}")
else:
    st.info("Please upload an image to begin.")
    if result:
    st.subheader("📌 Extracted Text:")
    extracted_lines = []
    for bbox, text, confidence in result:
        st.markdown(f"- **{text}** (Confidence: `{confidence:.2f}`)")
        extracted_lines.append({"Text": text, "Confidence": confidence})

    # Convert to DataFrame
    df = pd.DataFrame(extracted_lines)

    # Export as CSV
    csv = df.to_csv(index=False).encode("utf-8")
    st.download_button("📥 Download CSV", csv, "ocr_output.csv", "text/csv")

    # Export as Excel
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="OCR Text")
    st.download_button("📥 Download Excel", excel_buffer.getvalue(), "ocr_output.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # Export as DOCX
    doc = Document()
    doc.add_heading("Extracted OCR Text", level=1)
    for row in extracted_lines:
        doc.add_paragraph(f"{row['Text']} (Confidence: {row['Confidence']:.2f})")
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    st.download_button("📥 Download DOCX", doc_buffer.getvalue(), "ocr_output.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

