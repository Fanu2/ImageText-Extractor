import tkinter as tk
from tkinter import filedialog, messagebox, Text
from tkinter.ttk import Button
from PIL import Image, ImageTk, ImageEnhance, ImageFilter
import pytesseract
import pandas as pd
from io import BytesIO
from docx import Document

# -----------------------------
# Simple Helpers
# -----------------------------
def preprocess_image(img, gray=False, sharpen=False, threshold=False):
    processed = img.copy()

    if gray:
        processed = processed.convert("L")

    if threshold:
        bw = processed.convert("L")
        bw = bw.point(lambda x: 255 if x > 128 else 0)
        processed = bw

    if sharpen:
        processed = processed.filter(ImageFilter.SHARPEN)

    return processed


# -----------------------------
# GUI App
# -----------------------------
class SimpleOCRApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Simple OCR App")
        self.root.geometry("900x600")
        self.root.resizable(False, False)

        self.image = None
        self.processed = None

        # ------- TOP Buttons -------
        tk.Button(root, text="Open Image", command=self.load_image, width=20).pack(pady=5)
        
        tk.Button(root, text="Run OCR", command=self.run_ocr, width=20).pack(pady=5)

        # ------- Preprocessing options -------
        opt = tk.Frame(root)
        opt.pack(pady=5)

        self.var_gray = tk.IntVar()
        self.var_sharp = tk.IntVar()
        self.var_bw = tk.IntVar()

        tk.Checkbutton(opt, text="Grayscale", variable=self.var_gray).pack(side="left", padx=10)
        tk.Checkbutton(opt, text="Sharpen", variable=self.var_sharp).pack(side="left", padx=10)
        tk.Checkbutton(opt, text="Black/White", variable=self.var_bw).pack(side="left", padx=10)

        # ------- Two panels: Image + Text -------
        pan = tk.Frame(root)
        pan.pack(fill="both", expand=True)

        # Left panel image
        self.lbl_img = tk.Label(pan, text="(No image)", width=50, height=20, relief="sunken")
        self.lbl_img.pack(side="left", padx=10, pady=10)

        # Right panel text
        self.txt = Text(pan, wrap="word", width=50, height=30)
        self.txt.pack(side="right", padx=10, pady=10)

        # ------- Save buttons -------
        save_frame = tk.Frame(root)
        save_frame.pack(pady=5)

        Button(save_frame, text="Save as TXT", command=self.save_txt).pack(side="left", padx=10)
        Button(save_frame, text="Save as CSV", command=self.save_csv).pack(side="left", padx=10)
        Button(save_frame, text="Save as DOCX", command=self.save_docx).pack(side="left", padx=10)
        Button(save_frame, text="Save as Excel", command=self.save_excel).pack(side="left", padx=10)

    # -----------------------------
    # Load Image
    # -----------------------------
    def load_image(self):
        path = filedialog.askopenfilename(filetypes=[
            ("Images", "*.jpg *.jpeg *.png *.webp"),
            ("All files", "*.*")
        ])
        if not path:
            return
        
        try:
            self.image = Image.open(path).convert("RGB")
            self._display_image(self.image)
            self.txt.delete("1.0", tk.END)
        except Exception as e:
            messagebox.showerror("Error Loading Image", str(e))

    def _display_image(self, img):
        w, h = img.size
        scale = min(400 / w, 400 / h, 1)
        img_resized = img.resize((int(w*scale), int(h*scale)))
        tk_img = ImageTk.PhotoImage(img_resized)
        self.lbl_img.config(image=tk_img)
        self.lbl_img.image = tk_img

    # -----------------------------
    # Run OCR
    # -----------------------------
    def run_ocr(self):
        if not self.image:
            messagebox.showerror("Error", "Please load an image first.")
            return

        self.processed = preprocess_image(
            self.image,
            gray=self.var_gray.get(),
            sharpen=self.var_sharp.get(),
            threshold=self.var_bw.get()
        )

        try:
            text = pytesseract.image_to_string(self.processed)
            self.txt.delete("1.0", tk.END)
            self.txt.insert(tk.END, text)
        except Exception as e:
            messagebox.showerror("OCR Error", str(e))

    # -----------------------------
    # Save Functions
    # -----------------------------
    def save_txt(self):
        text = self.txt.get("1.0", tk.END).strip()
        if not text:
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt")
        if not path:
            return
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)

    def save_csv(self):
        text = self.txt.get("1.0", tk.END).strip()
        if not text:
            return
        df = pd.DataFrame({"Text": text.split("\n")})
        path = filedialog.asksaveasfilename(defaultextension=".csv")
        if not path:
            return
        df.to_csv(path, index=False)

    def save_docx(self):
        text = self.txt.get("1.0", tk.END).strip()
        if not text:
            return
        doc = Document()
        doc.add_heading("OCR Output", level=1)
        for line in text.split("\n"):
            doc.add_paragraph(line)
        path = filedialog.asksaveasfilename(defaultextension=".docx")
        if not path:
            return
        doc.save(path)

    def save_excel(self):
        text = self.txt.get("1.0", tk.END).strip()
        if not text:
            return
        df = pd.DataFrame({"Text": text.split("\n")})
        path = filedialog.asksaveasfilename(defaultextension=".xlsx")
        if not path:
            return
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="OCR")

# -----------------------------
# Main
# -----------------------------
def main():
    root = tk.Tk()
    SimpleOCRApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
